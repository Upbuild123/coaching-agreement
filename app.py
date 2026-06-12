"""
Upbuild Contract Automation — Streamlit App

Environment variables (set in .env locally, or Streamlit Cloud secrets):
  GMAIL_USER      — sending Gmail address
  GMAIL_PASSWORD  — 16-char Gmail App Password
  NOTIFY_EMAIL    — address to receive the agreement (defaults to GMAIL_USER)
"""

import io
import os
import shutil
import smtplib
import tempfile
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

import docx
import streamlit as st
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

import fee_builder as fb

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

COACHES = [
    "Rasanath Das",
    "Hari Prasada Das",
    "Michael Sloyer",
    "Tzipi Weiss",
    "Vipin Goyal",
]

COACH_EMAILS = {
    "Rasanath Das":    "rasanath@upbuild.com",
    "Hari Prasada Das": "hari@upbuild.com",
    "Tzipi Weiss":     "tzipi@upbuild.com",
    "Vipin Goyal":     "vipin@upbuild.com",
}

TEMPLATE = Path(__file__).parent / "Upbuild Coaching Agreement - Sample.docx"
SERVICES_TEMPLATE = Path(__file__).parent / "Upbuild Services Agreement - Template.docx"


# ---------------------------------------------------------------------------
# Env loader (local dev only; Streamlit Cloud uses st.secrets)
# ---------------------------------------------------------------------------

def load_env():
    env_file = Path(__file__).parent / ".env"
    if env_file.exists():
        for line in env_file.read_text().splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip())


def get_secret(key: str, default: str = "") -> str:
    """Read from Streamlit secrets first, fall back to env vars."""
    try:
        return st.secrets[key]
    except (KeyError, FileNotFoundError):
        return os.environ.get(key, default)


# ---------------------------------------------------------------------------
# Agreement generation
# ---------------------------------------------------------------------------

def replace_in_paragraph(para, old: str, new: str):
    if old == new or old not in para.text:
        return
    runs = para.runs
    full = "".join(r.text for r in runs)
    positions = []
    idx = 0
    while True:
        i = full.find(old, idx)
        if i == -1:
            break
        positions.append(i)
        idx = i + len(old)
    for occ_start in reversed(positions):
        occ_end = occ_start + len(old)
        pos, first = 0, False
        for run in runs:
            rs, re_ = pos, pos + len(run.text)
            if re_ > occ_start and rs < occ_end:
                pre  = run.text[:max(0, occ_start - rs)]
                post = run.text[max(0, occ_end - rs):]
                run.text = (pre + new + post) if not first else (pre + post)
                first = True
            pos = re_


def build_agreement_bytes(first_name: str, last_name: str, coach: str, rate: str) -> bytes:
    client_full = f"{first_name} {last_name}"
    rate = f"{int(rate):,}"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    shutil.copy(TEMPLATE, tmp_path)
    document = docx.Document(tmp_path)

    replacements = [
        ("[Client Full Name]",  client_full),
        ("[Client First Name]", first_name),
        ("[Coach Name]",        coach),
        ("[Rate]",              rate),
    ]

    for para in document.paragraphs:
        for old, new in replacements:
            replace_in_paragraph(para, old, new)

    buf = io.BytesIO()
    document.save(buf)
    tmp_path.unlink(missing_ok=True)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Services Agreement generation
# ---------------------------------------------------------------------------

def _insert_lines_after(paragraph, lines):
    """Insert each string in `lines` as a new paragraph after `paragraph`.

    Lines starting with "• " are rendered as bullet points.
    """
    anchor = paragraph._p
    for line in lines:
        new_p = OxmlElement("w:p")
        anchor.addnext(new_p)
        anchor = new_p
        new_para = Paragraph(new_p, paragraph._parent)
        if line.startswith("• "):
            new_para.style = "List Bullet"
            new_para.add_run(line[2:])
        else:
            new_para.add_run(line)


def _remove_paragraph(paragraph):
    p = paragraph._p
    p.getparent().remove(p)


def build_services_agreement_bytes(client_company: str, client_signer: str,
                                     fee_lines: list, cancellation_lines: list,
                                     invoicing_line: str) -> bytes:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    shutil.copy(SERVICES_TEMPLATE, tmp_path)
    document = docx.Document(tmp_path)

    replacements = [
        ("[Client Company Name]", client_company),
        ("[Client Signer Name]",  client_signer),
    ]

    fee_para = None
    cancellation_para = None
    invoicing_para = None
    for para in document.paragraphs:
        for old, new in replacements:
            replace_in_paragraph(para, old, new)
        if "[Fee Section]" in para.text:
            fee_para = para
        elif "[Cancellation Policy]" in para.text:
            cancellation_para = para
        elif para.text.strip().startswith("Invoices will be issued"):
            invoicing_para = para

    if invoicing_para is not None and invoicing_line:
        for run in invoicing_para.runs[1:]:
            run.text = ""
        if invoicing_para.runs:
            invoicing_para.runs[0].text = invoicing_line
        else:
            invoicing_para.add_run(invoicing_line)

    if cancellation_para is not None:
        if cancellation_lines:
            _insert_lines_after(cancellation_para, cancellation_lines)
        _remove_paragraph(cancellation_para)

    if fee_para is not None:
        _insert_lines_after(fee_para, fee_lines)
        _remove_paragraph(fee_para)

    buf = io.BytesIO()
    document.save(buf)
    tmp_path.unlink(missing_ok=True)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Email
# ---------------------------------------------------------------------------

def send_email(first_name: str, last_name: str, client_email: str,
               coach: str, rate: str, doc_bytes: bytes):
    gmail_user = get_secret("GMAIL_USER")
    gmail_pass = get_secret("GMAIL_PASSWORD")
    notify_to  = get_secret("NOTIFY_EMAIL") or gmail_user

    recipients = [notify_to]
    if coach in COACH_EMAILS:
        recipients.append(COACH_EMAILS[coach])

    client_full = f"{first_name} {last_name}"
    filename    = f"{client_full} Upbuild Agreement.docx"

    msg = MIMEMultipart()
    msg["From"]    = gmail_user
    msg["To"]      = ", ".join(recipients)
    msg["Subject"] = f"New Coaching Agreement — {client_full}"

    body = (
        f"A new coaching agreement has been generated.\n\n"
        f"  Client:  {client_full}\n"
        f"  Email:   {client_email}\n"
        f"  Coach:   {coach}\n"
        f"  Rate:    ${rate}/session\n\n"
        f"The agreement is attached."
    )
    msg.attach(MIMEText(body, "plain"))

    attachment = MIMEApplication(doc_bytes, Name=filename)
    attachment["Content-Disposition"] = f'attachment; filename="{filename}"'
    msg.attach(attachment)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(gmail_user, gmail_pass)
        server.sendmail(gmail_user, recipients, msg.as_string())


def send_services_agreement_email(client_company: str, client_signer: str,
                                    client_email: str, doc_bytes: bytes):
    gmail_user = get_secret("GMAIL_USER")
    gmail_pass = get_secret("GMAIL_PASSWORD")
    notify_to  = get_secret("NOTIFY_EMAIL") or gmail_user

    recipients = [notify_to]
    filename   = f"{client_company} Upbuild Services Agreement.docx"

    msg = MIMEMultipart()
    msg["From"]    = gmail_user
    msg["To"]      = ", ".join(recipients)
    msg["Subject"] = f"New Services Agreement — {client_company}"

    body = (
        f"A new services agreement has been generated.\n\n"
        f"  Client company: {client_company}\n"
        f"  Signer:         {client_signer}\n"
        f"  Email:          {client_email}\n\n"
        f"The agreement is attached."
    )
    msg.attach(MIMEText(body, "plain"))

    attachment = MIMEApplication(doc_bytes, Name=filename)
    attachment["Content-Disposition"] = f'attachment; filename="{filename}"'
    msg.attach(attachment)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(gmail_user, gmail_pass)
        server.sendmail(gmail_user, recipients, msg.as_string())


# ---------------------------------------------------------------------------
# Streamlit UI
# ---------------------------------------------------------------------------

load_env()

st.set_page_config(page_title="Upbuild | New Coaching Agreement", page_icon="📄", layout="centered")

st.markdown("""
<style>
    header[data-testid="stHeader"] {
        background: #14141f;
        height: 4rem;
    }
    .stApp {
        background: #f4f5f7;
    }
    .block-container {
        max-width: 640px;
        padding-top: 2rem !important;
    }
    #upbuild-topbar {
        background: #14141f;
        color: #ffffff;
        margin: -5rem -1px 2.5rem -1px;
        padding: 1.1rem 2rem;
        border-radius: 0 0 12px 12px;
        font-size: 1.15rem;
        font-weight: 700;
        letter-spacing: 0.04em;
        display: flex;
        align-items: center;
        gap: 0.5rem;
        position: relative;
        z-index: 999;
        height: 4rem;
    }
    #upbuild-heading h1 {
        font-size: 1.8rem;
        font-weight: 800;
        color: #1a1a2e;
        margin-bottom: 0.35rem;
    }
    #upbuild-heading p {
        color: #6b7280;
        font-size: 0.95rem;
        margin-bottom: 1.75rem;
    }
    div[data-testid="stForm"] {
        background: #ffffff;
        border-radius: 14px;
        padding: 2rem 2.25rem 2.25rem;
        box-shadow: 0 1px 3px rgba(0,0,0,0.06), 0 1px 12px rgba(0,0,0,0.04);
        border: 1px solid #eaeaf0;
    }
    .stTextInput input, .stSelectbox div[data-baseweb="select"] > div {
        border-radius: 8px !important;
        border-color: #e0e0e8 !important;
    }
    .stTextInput input:focus, .stSelectbox div[data-baseweb="select"]:focus-within > div {
        border-color: #4c2889 !important;
        box-shadow: 0 0 0 1px #4c2889 !important;
    }
    .stButton button,
    .stFormSubmitButton button,
    button[kind="primary"],
    button[kind="primaryFormSubmit"],
    div[data-testid="stFormSubmitButton"] button,
    button[data-testid="stBaseButton-primary"],
    button[data-testid="stBaseButton-primaryFormSubmit"] {
        background-color: #4c2889 !important;
        color: #fff !important;
        border-radius: 8px !important;
        border: none !important;
        font-weight: 600 !important;
        padding: 0.65rem 1.5rem !important;
        transition: background-color 0.15s ease !important;
    }
    .stButton button:hover,
    .stFormSubmitButton button:hover,
    button[kind="primary"]:hover,
    button[kind="primaryFormSubmit"]:hover,
    div[data-testid="stFormSubmitButton"] button:hover,
    button[data-testid="stBaseButton-primary"]:hover,
    button[data-testid="stBaseButton-primaryFormSubmit"]:hover {
        background-color: #3a1f6b !important;
        color: #fff !important;
    }
    label {
        font-weight: 600 !important;
        color: #374151 !important;
        font-size: 0.9rem !important;
    }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div id="upbuild-topbar">🥾 UPBUILD</div>
<div id="upbuild-heading">
    <h1>New Agreement</h1>
    <p>Fill in the details below to generate and send the agreement.</p>
</div>
""", unsafe_allow_html=True)

agreement_type = st.selectbox(
    "Agreement Type",
    ["Individual Coaching Agreement", "Services Agreement"],
    index=0,
)

if agreement_type == "Individual Coaching Agreement":
    with st.form("agreement_form"):
        col1, col2 = st.columns(2)
        first_name = col1.text_input("Client first name", placeholder="Jane")
        last_name  = col2.text_input("Client last name",  placeholder="Smith")

        client_email = st.text_input("Client email", placeholder="jane@example.com",
                                      help="Not added to the agreement — sent to you for follow-up.")

        coach = st.selectbox("Coach", options=COACHES)

        rate = st.text_input("Session rate ($)", placeholder="600",
                              help="Numbers only, e.g. 600")

        st.write("")
        submitted = st.form_submit_button("Generate & Send Agreement", type="primary", use_container_width=True)

    if submitted:
        errors = []
        if not first_name.strip():
            errors.append("First name is required.")
        if not last_name.strip():
            errors.append("Last name is required.")
        if not client_email.strip() or "@" not in client_email:
            errors.append("A valid client email is required.")
        if not rate.strip().isdigit():
            errors.append("Rate must be a whole number (e.g. 600).")

        if errors:
            for e in errors:
                st.error(e)
        else:
            with st.spinner("Generating agreement and sending email…"):
                try:
                    doc_bytes = build_agreement_bytes(
                        first_name.strip(), last_name.strip(), coach, rate.strip()
                    )
                    send_email(
                        first_name.strip(), last_name.strip(),
                        client_email.strip(), coach, rate.strip(),
                        doc_bytes,
                    )
                    st.success(
                        f"Agreement for {first_name.strip()} {last_name.strip()} "
                        f"generated and sent to you!"
                    )
                except Exception as exc:
                    st.error(f"Something went wrong: {exc}")

else:
    # -----------------------------------------------------------------------
    # Services Agreement
    # -----------------------------------------------------------------------

    st.subheader("Engagement Details")
    col1, col2 = st.columns(2)
    client_company = col1.text_input("Client company name", placeholder="Acme Corp", key="sa_client_company")
    client_signer  = col2.text_input("Client signer name", placeholder="Jane Smith", key="sa_client_signer")
    client_email   = st.text_input("Client contact email", placeholder="jane@example.com",
                                    help="Not added to the agreement — sent to you for follow-up.",
                                    key="sa_client_email")

    st.divider()
    st.subheader("Step 1: Engagement Type")
    engagement_type = st.selectbox("Engagement type", fb.ENGAGEMENT_TYPES, key="sa_engagement_type")

    if st.button("Apply defaults for this engagement type"):
        defaults = fb.default_state(engagement_type)

        st.session_state["sa_n_fee_items"] = len(defaults["fee_items"])
        for i, item in enumerate(defaults["fee_items"]):
            st.session_state[f"sa_fee_{i}_name"] = item["name"]
            st.session_state[f"sa_fee_{i}_pricing_type"] = item["pricing_type"]
            st.session_state[f"sa_fee_{i}_rate"] = item["rate"]
            st.session_state[f"sa_fee_{i}_notes"] = item["notes"]

        st.session_state["sa_fee_structure_mode"] = defaults["fee_structure_mode"]
        st.session_state["sa_estimate_amount"] = defaults["estimate_amount"]
        st.session_state["sa_estimate_disclaimer"] = defaults["estimate_disclaimer"]
        st.session_state["sa_include_additional"] = defaults["include_additional_services"]
        st.session_state["sa_additional_text"] = defaults["additional_services_text"]
        st.session_state["sa_include_expenses"] = defaults["include_expenses"]
        st.session_state["sa_expense_categories"] = defaults["expense_categories"]
        st.session_state["sa_expense_preapproval"] = defaults["expense_preapproval"]
        st.session_state["sa_travel_class"] = defaults["travel_class"]
        st.session_state["sa_travel_cap"] = defaults["travel_cap"]
        st.session_state["sa_expense_approver"] = defaults["expense_approver"]
        st.session_state["sa_invoice_frequency"] = defaults["invoice_frequency"]
        st.session_state["sa_payment_terms"] = defaults["payment_terms"]

        st.session_state["sa_n_cancel"] = len(defaults["cancellation_policies"])
        for i, pol in enumerate(defaults["cancellation_policies"]):
            st.session_state[f"sa_cancel_{i}_policy_type"] = pol["policy_type"]
            st.session_state[f"sa_cancel_{i}_service_type"] = pol["service_type"]
            st.session_state[f"sa_cancel_{i}_notice_period"] = pol["notice_period"]
            st.session_state[f"sa_cancel_{i}_charge_text"] = pol["charge_text"]
            st.session_state[f"sa_cancel_{i}_additional_text"] = pol["additional_text"]

        st.session_state["sa_sync_text"] = True
        st.rerun()

    st.divider()
    st.subheader("Step 2: Fee Items")

    st.session_state.setdefault("sa_n_fee_items", 1)
    n_fee_items = st.number_input("Number of fee items", min_value=1, max_value=10,
                                   step=1, key="sa_n_fee_items")

    fee_items = []
    for i in range(int(n_fee_items)):
        with st.expander(f"Fee Item {i + 1}", expanded=True):
            c1, c2 = st.columns(2)
            st.session_state.setdefault(f"sa_fee_{i}_name", "")
            name = c1.text_input("Name (as it appears in the agreement)",
                                  placeholder="e.g. 1:1 coaching", key=f"sa_fee_{i}_name")

            st.session_state.setdefault(f"sa_fee_{i}_pricing_type", fb.PRICING_TYPES[0])
            pricing_type = c2.selectbox("Pricing type", fb.PRICING_TYPES, key=f"sa_fee_{i}_pricing_type")

            c3, c4 = st.columns(2)
            st.session_state.setdefault(f"sa_fee_{i}_rate", "")
            rate_label = "Custom fee language" if pricing_type == "Custom Text" else "Rate ($)"
            rate_placeholder = "e.g. Fees will be billed at cost" if pricing_type == "Custom Text" else "650"
            rate = c3.text_input(rate_label, placeholder=rate_placeholder, key=f"sa_fee_{i}_rate")

            st.session_state.setdefault(f"sa_fee_{i}_notes", "")
            notes = c4.text_input("Notes (optional)", key=f"sa_fee_{i}_notes")

            fee_items.append({"name": name, "pricing_type": pricing_type, "rate": rate, "notes": notes})

    st.divider()
    st.subheader("Step 3: Fee Structure")
    st.session_state.setdefault("sa_fee_structure_mode", fb.FEE_STRUCTURE_MODES[0])
    fee_structure_mode = st.radio("Fee structure", fb.FEE_STRUCTURE_MODES,
                                   key="sa_fee_structure_mode", horizontal=True)

    estimate_amount = ""
    estimate_disclaimer = True
    if fee_structure_mode == "Estimated Annual Investment":
        ec1, ec2 = st.columns([2, 1])
        st.session_state.setdefault("sa_estimate_amount", "")
        estimate_amount = ec1.text_input("Estimated annual investment ($)",
                                          placeholder="150,000", key="sa_estimate_amount")
        st.session_state.setdefault("sa_estimate_disclaimer", True)
        estimate_disclaimer = ec2.checkbox("Include disclaimer", key="sa_estimate_disclaimer")

    st.divider()
    st.subheader("Step 4: Additional Services")
    st.session_state.setdefault("sa_include_additional", True)
    include_additional = st.checkbox("Include additional services language", key="sa_include_additional")

    additional_text = ""
    if include_additional:
        st.session_state.setdefault(
            "sa_additional_text",
            "Fees for any additional services will depend on the scope of the services "
            "and will be discussed and agreed upon in advance.",
        )
        additional_text = st.text_area("Additional services language", key="sa_additional_text", height=80)

    st.divider()
    st.subheader("Step 5: Expenses")
    st.session_state.setdefault("sa_include_expenses", False)
    include_expenses = st.checkbox("Include expense reimbursement", key="sa_include_expenses")

    expense_config = {"include_expenses": include_expenses}
    if include_expenses:
        st.session_state.setdefault("sa_expense_categories", ["Travel", "Lodging", "Meals"])
        expense_categories = st.multiselect("Reimbursable expenses", fb.EXPENSE_CATEGORIES,
                                             key="sa_expense_categories")
        expense_config["expense_categories"] = expense_categories

        if "Custom" in expense_categories:
            expense_config["expense_category_custom"] = st.text_input(
                "Custom expense category", key="sa_expense_category_custom")

        st.session_state.setdefault("sa_expense_preapproval", True)
        expense_config["expense_preapproval"] = st.checkbox("Pre-approval required?", key="sa_expense_preapproval")

        if "Travel" in expense_categories:
            tc1, tc2, tc3 = st.columns(3)
            st.session_state.setdefault("sa_travel_class", fb.TRAVEL_CLASSES[2])
            travel_class = tc1.selectbox("Travel class", fb.TRAVEL_CLASSES, key="sa_travel_class")
            expense_config["travel_class"] = travel_class
            if travel_class == "Custom":
                expense_config["travel_class_custom"] = tc1.text_input(
                    "Custom travel class", key="sa_travel_class_custom")

            st.session_state.setdefault("sa_travel_cap", "5,000")
            expense_config["travel_cap"] = tc2.text_input("Travel cap ($)", key="sa_travel_cap")

            st.session_state.setdefault("sa_expense_approver", "the Client")
            expense_config["expense_approver"] = tc3.text_input("Approver (name or role)", key="sa_expense_approver")

    st.divider()
    st.subheader("Step 6: Invoicing")
    ic1, ic2 = st.columns(2)
    st.session_state.setdefault("sa_invoice_frequency", fb.INVOICE_FREQUENCIES[0])
    invoice_frequency = ic1.selectbox("Invoice frequency", fb.INVOICE_FREQUENCIES, key="sa_invoice_frequency")
    invoicing_config = {"invoice_frequency": invoice_frequency}
    if invoice_frequency == "Custom":
        invoicing_config["invoice_frequency_custom"] = st.text_input(
            "Custom invoicing language", key="sa_invoice_frequency_custom")

    st.session_state.setdefault("sa_payment_terms", fb.PAYMENT_TERMS[0])
    payment_terms = ic2.selectbox("Payment terms", fb.PAYMENT_TERMS, key="sa_payment_terms")
    invoicing_config["payment_terms"] = payment_terms
    if payment_terms == "Custom":
        invoicing_config["payment_terms_custom"] = st.text_input(
            "Custom payment terms", key="sa_payment_terms_custom")

    st.divider()
    st.subheader("Step 7: Cancellation Policy")
    st.session_state.setdefault("sa_n_cancel", 1)
    n_cancel = st.number_input("Number of cancellation policies", min_value=1, max_value=5,
                                step=1, key="sa_n_cancel")

    cancellation_policies = []
    for i in range(int(n_cancel)):
        with st.expander(f"Cancellation Policy {i + 1}", expanded=True):
            st.session_state.setdefault(f"sa_cancel_{i}_policy_type", fb.CANCELLATION_TYPES[0])
            policy_type = st.selectbox("Policy type", fb.CANCELLATION_TYPES,
                                        key=f"sa_cancel_{i}_policy_type")

            st.session_state.setdefault(f"sa_cancel_{i}_service_type", "")
            service_type = st.text_input("Service type / label (e.g. 'individual coaching session')",
                                          key=f"sa_cancel_{i}_service_type")

            pc1, pc2 = st.columns(2)
            st.session_state.setdefault(f"sa_cancel_{i}_notice_period",
                                         "10" if policy_type == "Workshop / Offsite" else "2")
            notice_period = pc1.text_input("Notice period (business days)",
                                            key=f"sa_cancel_{i}_notice_period")

            st.session_state.setdefault(f"sa_cancel_{i}_charge_text",
                                         "in full" if policy_type == "Workshop / Offsite" else "the full session fee")
            charge_text = pc2.text_input("Charge if canceled late", key=f"sa_cancel_{i}_charge_text")

            st.session_state.setdefault(f"sa_cancel_{i}_additional_text", "")
            additional_text_pol = st.text_input("Additional language (optional)",
                                                  key=f"sa_cancel_{i}_additional_text")

            cancellation_policies.append({
                "policy_type": policy_type,
                "service_type": service_type,
                "notice_period": notice_period,
                "charge_text": charge_text,
                "additional_text": additional_text_pol,
            })

    # -------------------------------------------------------------------
    # Build generated text + live preview / manual edit
    # -------------------------------------------------------------------

    state = {
        "fee_items": fee_items,
        "fee_structure_mode": fee_structure_mode,
        "estimate_amount": estimate_amount,
        "estimate_disclaimer": estimate_disclaimer,
        "include_additional_services": include_additional,
        "additional_services_text": additional_text,
        "cancellation_policies": cancellation_policies,
        **expense_config,
        **invoicing_config,
    }

    generated_fee_text = fb.lines_to_text(fb.build_fee_section_lines(state))
    generated_cancellation_text = fb.lines_to_text(fb.build_cancellation_section_lines(state))
    generated_invoicing_text = fb.build_invoicing_text(state)[0]

    if st.session_state.pop("sa_sync_text", False):
        st.session_state["sa_fee_text"] = generated_fee_text
        st.session_state["sa_cancellation_text"] = generated_cancellation_text
        st.session_state["sa_invoicing_text"] = generated_invoicing_text
    st.session_state.setdefault("sa_fee_text", generated_fee_text)
    st.session_state.setdefault("sa_cancellation_text", generated_cancellation_text)
    st.session_state.setdefault("sa_invoicing_text", generated_invoicing_text)

    st.divider()
    st.subheader("Step 8: Section 3 — Live Preview & Manual Edit")

    if st.button("Reset text to generated section"):
        st.session_state["sa_fee_text"] = generated_fee_text
        st.session_state["sa_cancellation_text"] = generated_cancellation_text
        st.session_state["sa_invoicing_text"] = generated_invoicing_text
        st.rerun()

    fee_text = st.text_area("Fees, Additional Services & Expenses", key="sa_fee_text", height=180)
    cancellation_text = st.text_area(
        "Cancellation Policy (lines starting with \"• \" become bullet points)",
        key="sa_cancellation_text", height=120,
    )
    invoicing_text = st.text_area("Invoicing & Payment", key="sa_invoicing_text", height=70)

    with st.expander("Preview", expanded=True):
        for section_text in (fee_text, cancellation_text, invoicing_text):
            for line in section_text.splitlines():
                if line.strip().startswith("• "):
                    st.markdown(f"- {line.strip()[2:]}")
                elif line.strip():
                    st.write(line.strip())

    st.divider()
    if st.button("Generate & Send Services Agreement", type="primary", use_container_width=True):
        errors = []
        if not client_company.strip():
            errors.append("Client company name is required.")
        if not client_signer.strip():
            errors.append("Client signer name is required.")
        if not client_email.strip() or "@" not in client_email:
            errors.append("A valid client contact email is required.")

        if errors:
            for e in errors:
                st.error(e)
        else:
            with st.spinner("Generating agreement and sending email…"):
                try:
                    fee_lines = fb.text_to_lines(fee_text)
                    cancellation_lines = fb.text_to_lines(cancellation_text)
                    invoicing_line = " ".join(fb.text_to_lines(invoicing_text))
                    doc_bytes = build_services_agreement_bytes(
                        client_company.strip(), client_signer.strip(),
                        fee_lines, cancellation_lines, invoicing_line,
                    )
                    send_services_agreement_email(
                        client_company.strip(), client_signer.strip(),
                        client_email.strip(), doc_bytes,
                    )
                    st.success(f"Services Agreement for {client_company.strip()} generated and sent to you!")
                except Exception as exc:
                    st.error(f"Something went wrong: {exc}")
