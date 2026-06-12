"""One-off script to generate 'Upbuild Services Agreement - Template.docx'.

Run once locally: python3 build_services_template.py
"""

import docx
from docx.shared import Pt
from pathlib import Path

OUT = Path(__file__).parent / "Upbuild Services Agreement - Template.docx"

doc = docx.Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h1(text):
    doc.add_heading(text, level=1)


def p(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


# Title
title = doc.add_heading("Upbuild Services Agreement", level=0)

p(
    "This Agreement is entered into by and between: [Client Company Name] (Client) and "
    "Upbuild (Company) and is effective as of the date the Client signs below. Through a "
    "relationship with one or more Upbuild coaches (each, a Coach), the Company agrees to "
    "provide coaching, consulting, and/or mediation services for individual employees of the "
    "Client (Client Employees) focusing on topics and goals to be defined by the Client and/or "
    "Client Employees, and the Coach. This Agreement supersedes all prior agreements between "
    "the parties regarding the subject matter herein."
)

h1("1. Description of Services.")
bullet(
    "Coaching. Coaching is a partnership between the Coach and Client Employees in a "
    "thought-provoking and creative process that inspires Client Employees to maximize "
    "personal and professional potential (Coaching)."
)
bullet(
    "360 Degree Reviews. A “360” is a structured process facilitated by the Coach "
    "for the Client and/or Client Employees to get feedback from team members, peers, and "
    "other stakeholders (360 Degree Reviews)."
)
bullet(
    "Mediation. Mediation is a structured and customized process facilitated by the Coach to "
    "help identify gaps in trust and rebuild it between or among Client Employees where it has "
    "eroded."
)
bullet(
    "Workshops. Workshops are group learning opportunities facilitated by the Coach designed "
    "to develop applied knowledge in a specific area, build trust, foster connection, and "
    "unlock potential in the Client’s organization."
)
bullet(
    "Consulting. Throughout the term of this Agreement, at times convenient to the Coach and "
    "the Client, and upon the Client’s request, the Coach may provide additional consulting "
    "and/or professional services based on the hourly rate below."
)

h1("2. Coach-Client Relationship")
bullet(
    "The Coach agrees to maintain the ethics and standards of behavior established by the "
    "International Coach Federation (coachfederation.org/ethics)."
)
bullet(
    "The Client acknowledges that the services described above are comprehensive processes "
    "that may involve different areas of life, including work, finances, health, "
    "relationships, and education."
)
bullet(
    "The Client agrees that deciding how to handle these issues, incorporating coaching "
    "principles into those areas, and implementing choices is the sole responsibility of "
    "Client Employees."
)
bullet(
    "The Client acknowledges that the services described above are not therapy and do not "
    "substitute for therapy, if needed. The Client also understands that these services do not "
    "involve the diagnosis, treatment, or cure of mental disorders or medical condition or "
    "disease, and that they are not a substitute for counseling, psychotherapy, substance abuse "
    "treatment, or other medical, mental, or other professional advice. It is the Client’s "
    "team members, employees, or stakeholders’ responsibility to seek independent "
    "professional guidance, as needed."
)

h1("3. Fees, Expenses, and Cancellation Policy.")
doc.add_paragraph("[FEES_SECTION]")

h1("4. Termination.")
p(
    "Either the Client or the Company may terminate this Agreement at any time upon 48 hours’ "
    "notice, provided in writing via email or text. All fees for Services rendered will become "
    "due and payable within 30 days of termination and the parties’ rights and obligations "
    "described in this Agreement will survive the termination of the relationship. The Client "
    "agrees to compensate the Company for all Services rendered through the date of "
    "termination."
)
p(
    "When a Client Employee is ready to complete coaching, the Company requests that the "
    "Client Employee schedule a final coaching session as a “completion session” for closure."
)

h1("5. Confidentiality.")
p(
    "Unless the parties expressly agree otherwise in writing, all materials and content used or "
    "created by the Company or the Coach in connection with the Services are considered the "
    "Company’s proprietary and confidential information (“Confidential Information”) and "
    "belong to the Company. The Client agrees to keep confidential, not disclose, and to "
    "protect the Confidential Information. Notwithstanding, Confidential Information does not "
    "include information that is generally known to the public (by no fault of the Client) or "
    "was independently developed by the Client and does not relate to the Services provided "
    "hereunder."
)
p(
    "Additionally, the Coach/Client relationship is one of confidence and trust. The Coach "
    "agrees to keep all information shared by the Client confidential, except when maintaining "
    "such confidentiality would violate the law or the International Coaching Federation (ICF) "
    "Code of Ethics. Notwithstanding the foregoing, the Coach may share Client information "
    "within the Company for training and support purposes, and with Client Employee’s "
    "consent, the Company may share Client Employee’s name and email address with the ICF "
    "for credentialing purposes."
)

h1("6. Independent Contractor.")
p(
    "At all times the Company and the Client shall be deemed to be in a business-to-business "
    "relationship. Nothing herein shall be construed to create or imply that there exists "
    "between the parties a partnership, joint venture or other combined business organization. "
    "The Coach holds no authority, express or implied, to commit, obligate, or make "
    "representations on behalf of the Client and shall make no representation to others to the "
    "contrary. Nothing herein is intended nor shall be construed for any purpose as creating "
    "the relation of employer and employee or agent and principal between the parties. Except "
    "as otherwise specified herein, the Coach retains the right to direct, control or supervise "
    "the details and means by which the Services are provided. The Coach shall not be eligible "
    "for, or participate in, any insurance, pension, workers’ compensation insurance, profit "
    "sharing or other plans established for the benefit of Client Employees."
)
p(
    "The Company shall be responsible for payment of all taxes arising out of the Coach’s "
    "activities in connection with this Agreement, including without limitation, federal and "
    "state income taxes, social security taxes, unemployment insurance taxes, and any other "
    "taxes or business license fees as required. The Client shall not be responsible for "
    "withholding any income or employment taxes whatsoever on behalf of the Coach."
)

h1("7. Limited Liability and Indemnification.")
p(
    "Except as outlined in this Agreement, neither the Company nor the Coach makes any "
    "guarantees with respect to the Services, and will not be liable to the Client for any "
    "indirect, consequential, or special damages. The Client and its team members, employees, "
    "and other stakeholders are responsible for his/her/their own physical, mental and "
    "emotional well-being, decisions, choices, actions and results, including without "
    "limitation the Client’s or the team member’s, employee’s, or other stakeholder’s "
    "decisions they make as a result of the Services. To be clear, neither the Coach nor the "
    "Company is responsible or liable for any actions or inaction, or for any direct or "
    "indirect result of any of the Services, including without limitation if a team member, "
    "employee, or other stakeholder decides to quit or is terminated as a direct or indirect "
    "result of the Company’s Services."
)
p(
    "ALL SERVICES ARE PROVIDED “AS IS” WITHOUT WARRANTY OF ANY KIND, EITHER EXPRESS OR "
    "IMPLIED, INCLUDING BUT NOT LIMITED TO THE IMPLIED WARRANTIES OF MERCHANTABILITY AND "
    "FITNESS FOR A PARTICULAR PURPOSE. THE ENTIRE RISK AS TO THE QUALITY AND PERFORMANCE OF THE "
    "SERVICES IS WITH THE CLIENT. The Client shall indemnify and hold harmless the Company and "
    "its affiliates, directors, officers, employees, partners, contractors, coaches, or agents, "
    "from and against any and all claims, actions, causes of action, demands, or liabilities of "
    "whatsoever kind and nature, including judgments, interest, reasonable attorneys’ fees, "
    "and all other costs, fees, expenses, and charges (collectively, “Claims”) to the contrary "
    "which are brought by any third party including but not limited to the Client and its "
    "employees, team members, partners or other stakeholders."
)
p(
    "IN NO EVENT SHALL THE COMPANY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, "
    "PUNITIVE OR CONSEQUENTIAL DAMAGES INCURRED BY THE CLIENT OR ANY THIRD PARTY, WHETHER IN "
    "CONTRACT, TORT, OR BASED UPON A WARRANTY, EVEN IF THE CLIENT OR ANY THIRD PARTY HAS BEEN "
    "ADVISED OF THE POSSIBILITY OF SUCH DAMAGES. THE TOTAL LIABILITY OF THE COMPANY TO THE "
    "CLIENT SHALL NOT EXCEED THE SUM OF THE FEES PAID TO THE COMPANY BY THE CLIENT HEREUNDER."
)

h1("8. Applicable Law and Attorneys’ Fees.")
p(
    "This Agreement shall be governed and construed in accordance with the laws of New York. "
    "In the event of any legal dispute concerning any controversy between the parties or "
    "otherwise arising out of or related to the subject matter of this Agreement, the "
    "prevailing party shall be entitled to recover from the losing party reasonable expenses, "
    "attorneys’ fees, pre and post-judgment interest, and all other costs, expenses and fees "
    "incurred in connection with the enforcement of the terms of this Agreement at all trial "
    "and appellate levels including demonstrating the existence of a breach, seeking and "
    "obtaining equitable relief, and any other enforcement efforts, from the non-prevailing "
    "party."
)

h1("9. Severability.")
p(
    "The parties intend this Agreement to be enforced to the fullest extent permissible under "
    "the laws and public policies. If any particular provision of this Agreement as "
    "adjudicated is invalid, prohibited or unenforceable for any reason, such provision shall "
    "be ineffective, without invalidating the remaining provisions of this Agreement or "
    "affecting the validity or enforceability of this Agreement."
)

p(
    "The Client’s authorized signature below indicates that the Client has read and agrees to "
    "abide by the terms and conditions of this Agreement. The Client affirms and agrees that no "
    "promises or agreements which are not herein expressed have been made to the Client in "
    "executing this Agreement."
)

doc.add_paragraph("")
p("Client")
p("Printed Name: [Client Signer Name], on behalf of [Client Company Name]")
p("Signature:")
doc.add_paragraph("")
p("[!Sign.1.F,client, ,]")

doc.add_paragraph("")
p("Company")
p("Printed Name: Ariel Weiss, on behalf of Upbuild")
p("Signature:")
doc.add_paragraph("")
p("[!Sign.2.F,Partner, ,]")

doc.save(OUT)
print(f"Saved {OUT}")
